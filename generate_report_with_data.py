from docx import Document
import os

class Report:
    OVERWRITE_PERM = TRUE
    def __init__(self,report_name,report_data,info_columns_dict):
        self.report_name = report_name
        self.report_data = report_data
        self.info_columns_dict = info_columns_dict
        self.doc = Document()
        self.gen_title()
        self.gen_info_paragraph()
        self.questions_indexs = list(range(3,len(self.report_data)))
        for i in self.questions_indexs:
            self.question_and_answer_writing(i)

    def gen_title(self):
        self.doc.add_heading(self.report_name)

    def gen_info_paragraph(self):
        self.doc.add_paragraph(f"""
姓名：{self.report_data[self.info_columns_dict['name']]}\t班级：{self.report_data[self.info_columns_dict['class']]}\t学号：{self.report_data[self.info_columns_dict['id']]}
""")

    def question_and_answer_writing(self,ind):
        self.doc.add_paragraph(self.report_data.index[ind])
        self.doc.add_paragraph(self.report_data[ind])

    def report_save(self):
        file_path = f'{self.report_data[self.info_columns_dict["id"]]}_{self.report_data[self.info_columns_dict["name"]]}.docx'
        if os.path.exists(file_path) & ~self.OVERWRITE_PERM:
            print(f'{file_path} already exist, no overwrites. this can be override in class coding')
            return None
        self.doc.save(file_path)
